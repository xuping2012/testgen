#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档解析模块
支持多种格式的文档处理
"""

from docx import Document
from PyPDF2 import PdfReader
import openpyxl
import os
import cv2
import pytesseract
import mistune

def parse_docx(filepath):
    """
    解析docx文档
    """
    try:
        doc = Document(filepath)
        text = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(' | '.join(row_text))
        
        return '\n'.join(text)
    except Exception as e:
        raise Exception(f"解析docx文档失败: {str(e)}")

def parse_pdf(filepath):
    """
    解析PDF文档
    """
    try:
        reader = PdfReader(filepath)
        text = []
        
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text.append(page_text.strip())
        
        return '\n'.join(text)
    except Exception as e:
        raise Exception(f"解析PDF文档失败: {str(e)}")

def parse_txt(filepath):
    """
    解析txt文档
    """
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        raise Exception(f"解析txt文档失败: {str(e)}")

def parse_image(filepath):
    """
    解析图片文档（OCR）
    """
    try:
        # 读取图片
        img = cv2.imread(filepath)
        if img is None:
            raise Exception("无法读取图片")
        
        # 转换为灰度图
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # OCR识别
        text = pytesseract.image_to_string(gray, lang='chi_sim')
        return text
    except Exception as e:
        raise Exception(f"解析图片文档失败: {str(e)}")

def parse_markdown(filepath):
    """
    解析markdown文档
    """
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # 直接返回文本内容，不使用mistune解析
        # 因为mistune 3.x版本API与之前不同
        return content
    except Exception as e:
        raise Exception(f"解析markdown文档失败: {str(e)}")

def parse_excel(filepath):
    """
    解析Excel文件
    """
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True)
        text = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text.append(f"## 工作表: {sheet_name}")
            
            for row in ws.iter_rows(values_only=True):
                # 过滤空行
                if any(cell is not None for cell in row):
                    # 将单元格用 | 分隔，模拟表格格式
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    text.append(row_text)
            text.append("")  # 空行分隔
        
        wb.close()
        return "\n".join(text)
    except Exception as e:
        raise Exception(f"解析Excel文件失败: {str(e)}")

def parse_document(filepath):
    """
    根据文件扩展名选择相应的解析方法
    """
    if not os.path.exists(filepath):
        raise Exception(f"文件不存在: {filepath}")
    
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.docx':
        return parse_docx(filepath)
    elif ext == '.pdf':
        return parse_pdf(filepath)
    elif ext == '.txt':
        return parse_txt(filepath)
    elif ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
        return parse_image(filepath)
    elif ext in ['.md', '.markdown']:
        return parse_markdown(filepath)
    elif ext in ['.xlsx', '.xls']:
        return parse_excel(filepath)
    else:
        raise Exception(f"不支持的文件格式: {ext}")

if __name__ == '__main__':
    # 测试解析功能
    test_file = 'Demo AI提效生成测试用例.docx'
    if os.path.exists(test_file):
        content = parse_document(test_file)
        print(f"解析成功，内容长度: {len(content)}")
        print("\n前500字符:")
        print(content[:500])
    else:
        print(f"测试文件不存在: {test_file}")
